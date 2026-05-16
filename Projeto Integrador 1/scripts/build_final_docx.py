"""Gera o relatorio final DOCX do Projeto Integrador 1."""

from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = ROOT / "modelo" / "Modelo de Projeto CDML.docx"
OUTPUT_PATH = ROOT / "outputs" / "reports" / "Predicao_Chuvas_Brasilia_Enzo_Innecco.docx"
BACKTEST_PATH = ROOT / "outputs" / "reports" / "rain_model_backtest.json"
RUN_REPORT_PATH = ROOT / "outputs" / "reports" / "daily_pipeline_run_report.json"
CITY_SUMMARY_PATH = ROOT / "outputs" / "reports" / "weather_city_comparison_summary.csv"
WORKING_TEMPLATE_PATH = ROOT / "outputs" / "reports" / "_template_working_copy.docx"

FIGURES = [
    (
        ROOT / "outputs" / "figures" / "eda_taxa_chuva_mensal_por_cidade.png",
        "Figura 1 - Taxa mensal de dias com chuva por cidade.",
    ),
    (
        ROOT / "outputs" / "figures" / "eda_precipitacao_media_mensal_por_cidade.png",
        "Figura 2 - Precipitação média diária por mês.",
    ),
    (
        ROOT / "outputs" / "figures" / "eda_temperatura_media_mensal_por_cidade.png",
        "Figura 3 - Temperatura média diária por mês.",
    ),
    (
        ROOT / "outputs" / "figures" / "eda_precipitacao_anual_por_cidade.png",
        "Figura 4 - Precipitação anual acumulada por cidade.",
    ),
]

PROJECT_NAME = "Predição Chuvas Brasília"
STUDENT_NAME = "Enzo Innecco"
STUDENT_ID = "22253256"
NUMBERING_CACHE: dict[int, dict[str, int]] = {}


def main() -> None:
    backtest = json.loads(BACKTEST_PATH.read_text(encoding="utf-8"))
    run_report = json.loads(RUN_REPORT_PATH.read_text(encoding="utf-8"))
    city_summary = read_city_summary(CITY_SUMMARY_PATH)

    document = load_document()
    fill_cover(document)
    remove_template_body(document)
    configure_styles(document)
    build_report_body(document, backtest, run_report, city_summary)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    WORKING_TEMPLATE_PATH.unlink(missing_ok=True)
    print(OUTPUT_PATH)


def load_document() -> Document:
    """Carrega o modelo quando disponivel e cai para documento novo se necessario."""
    source_path = WORKING_TEMPLATE_PATH if WORKING_TEMPLATE_PATH.exists() else TEMPLATE_PATH
    try:
        return Document(source_path)
    except PermissionError:
        return Document()


def fill_cover(document: Document) -> None:
    """Preenche capa herdada do modelo."""
    student_line_filled = False
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text == "MATRÍCULA      Nome do Aluno":
            if not student_line_filled:
                set_paragraph_text(paragraph, f"{STUDENT_ID}      {STUDENT_NAME}")
                student_line_filled = True
            else:
                set_paragraph_text(paragraph, "")
        elif text == "NOME DO PROJETO":
            set_paragraph_text(paragraph, PROJECT_NAME.upper())
        elif text == "Brasília, <mês/ ano>":
            set_paragraph_text(paragraph, "Brasília, maio/2026")

    if not document.paragraphs:
        add_centered_heading(document, "CENTRO UNIVERSITÁRIO DE BRASÍLIA (CEUB)", 1)
        add_centered_heading(document, "CURSO DE CIÊNCIA DE DADOS E MACHINE LEARNING", 2)
        document.add_paragraph("")
        add_centered_heading(document, f"{STUDENT_ID}      {STUDENT_NAME}", 2)
        document.add_paragraph("")
        add_centered_heading(document, PROJECT_NAME.upper(), 1)
        document.add_paragraph("")
        add_centered_heading(document, "Brasília, maio/2026", 2)


def remove_template_body(document: Document) -> None:
    """Remove conteudo instrucional do modelo a partir do sumario."""
    start_index = None
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.text.strip() == "SUMARIO":
            start_index = index
            break
    if start_index is None:
        return
    for paragraph in list(document.paragraphs[start_index:]):
        element = paragraph._element
        element.getparent().remove(element)


def configure_styles(document: Document) -> None:
    """Configura estilos consistentes para o relatorio."""
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(31, 77, 120)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def build_report_body(
    document: Document,
    backtest: dict[str, object],
    run_report: dict[str, object],
    city_summary: list[dict[str, str]],
) -> None:
    document.add_page_break()
    add_centered_heading(document, "SUMÁRIO", level=1)
    for item in summary_items():
        document.add_paragraph(item)

    document.add_page_break()
    add_heading(document, "1. Introdução", 1)
    add_heading(document, "1.1 Contextualização do problema", 2)
    add_paragraph(
        document,
        "Brasília apresenta sazonalidade climática intensa, com alternância entre "
        "períodos secos e chuvosos. A ocorrência de chuva no dia seguinte afeta "
        "decisões operacionais de deslocamento, atividades externas, planejamento "
        "de eventos, aulas práticas e uso de infraestrutura urbana.",
    )
    add_paragraph(
        document,
        "O projeto propõe uma plataforma robusta de dados para estimar a ocorrência "
        "de chuva no dia seguinte em Brasília, com coleta automatizada, validação "
        "por contrato, comparação regional, análise exploratória e avaliação "
        "temporal do modelo.",
    )

    add_heading(document, "1.2 Justificativa", 2)
    add_paragraph(
        document,
        "O tema permite demonstrar um ciclo completo de engenharia de dados e "
        "machine learning com rastreabilidade: ingestão por API, camadas de dados, "
        "contratos de qualidade, variáveis preditivas, comparação entre cidades, "
        "orquestração diária e métricas auditáveis.",
    )

    add_heading(document, "1.3 Objetivo geral", 2)
    add_paragraph(
        document,
        "Construir uma solução robusta para previsão de ocorrência de chuva no dia "
        "seguinte em Brasília, usando dados meteorológicos históricos da Open-Meteo "
        "e comparação exploratória com Goiânia e São Paulo.",
    )

    add_heading(document, "1.4 Objetivos específicos", 2)
    add_numbered_list(
        document,
        [
            "Coletar dados meteorológicos históricos por API pública.",
            "Centralizar parâmetros sensíveis e variáveis operacionais em arquivo .env.",
            "Validar dados brutos e processados por contratos declarativos.",
            "Construir base supervisionada para previsão do dia seguinte.",
            "Gerar análise exploratória com comparação entre Brasília, Goiânia e São Paulo.",
            "Treinar e avaliar modelo explicável com validação temporal.",
            "Orquestrar atualização diária às 05:00 por DAG do Airflow.",
            "Organizar anexos técnicos para auditoria e apresentação acadêmica.",
        ],
    )

    add_heading(document, "2. Metodologia Design Science Research (DSR)", 1)
    add_heading(document, "2.1 Problema relevante", 2)
    add_paragraph(
        document,
        "A previsão de chuva em horizonte diário é relevante para planejamento "
        "operacional. O artefato desenvolvido combina utilidade prática, execução "
        "reprodutível e mecanismos de controle de qualidade.",
    )

    add_heading(document, "2.2 Construção do artefato", 2)
    add_bullet_list(
        document,
        [
            "Estrutura de projeto com dados brutos, dados processados, modelos, relatórios e figuras.",
            "Arquivo .env para URL, chave opcional de API, período, cidade principal e agendamento.",
            "Coletor multi-cidade para Brasília, Goiânia e São Paulo.",
            "Contratos de dados com validação de colunas, tipos e chave composta cidade-data.",
            "Pipeline operacional executável por script local e por Airflow.",
            "Gráficos de análise exploratória e resumo comparativo por cidade.",
            "Modelo probabilístico explicável para previsão de chuva em Brasília.",
        ],
    )

    add_heading(document, "2.3 Avaliação", 2)
    add_paragraph(
        document,
        "A avaliação utiliza testes automatizados, validação estrutural das bases, "
        "auditoria dos artefatos e holdout temporal. O treinamento usa histórico "
        "até 2024 e a validação considera dados posteriores a 2025-01-01.",
    )

    add_heading(document, "2.4 Construção científica", 2)
    add_paragraph(
        document,
        "A contribuição está na construção de um artefato de dados governado por "
        "contratos, com comparação regional e execução diária. O projeto evidencia "
        "como decisões técnicas fortalecem a confiabilidade do resultado analítico.",
    )

    add_heading(document, "2.5 Comunicação", 2)
    add_paragraph(
        document,
        "Os resultados são comunicados por relatório final, anexos técnicos, gráficos, "
        "arquivos de configuração, DAG do Airflow, código-fonte e relatórios JSON/CSV.",
    )

    add_heading(document, "3. Engenharia e Análise de Dados", 1)
    add_heading(document, "3.1 Fonte dos dados", 2)
    fetch = run_report["fetch"]
    add_table(
        document,
        ["Item", "Descrição"],
        [
            ["Fonte", "Open-Meteo Historical Weather API"],
            ["Cidades", "Brasília (DF), Goiânia (GO) e São Paulo (SP)"],
            ["Período coletado", f"{fetch['start_date']} a {fetch['end_date']}"],
            ["Registros coletados", str(fetch["row_count"])],
            ["Granularidade", "Diária"],
            ["Formato", "JSON na API; CSV no armazenamento local"],
            ["Chave de API", "Não obrigatória para uso não comercial; campo preparado no .env."],
        ],
        [2.0, 4.5],
    )

    add_heading(document, "3.2 Pré-processamento e engenharia de features", 2)
    add_paragraph(
        document,
        "O pipeline transforma observações diárias de Brasília em exemplos "
        "supervisionados. Cada linha usa variáveis do dia atual para prever a "
        "ocorrência de chuva no dia seguinte.",
    )
    add_table(
        document,
        ["Feature", "Descrição"],
        [
            ["month", "Mês da observação, usado para capturar sazonalidade."],
            ["temperature_range", "Diferença entre temperatura máxima e mínima."],
            ["rain_today", "Indicador binário de chuva no dia atual."],
            ["target_rain_tomorrow", "Alvo binário: 1 quando a precipitação do dia seguinte é >= 1 mm."],
        ],
        [2.0, 4.5],
    )

    add_heading(document, "3.3 Análise exploratória e comparação regional", 2)
    add_paragraph(
        document,
        "A análise exploratória compara Brasília com Goiânia e São Paulo para "
        "contextualizar o comportamento climático da cidade-alvo. O comparativo "
        "mostra diferença de taxa de chuva, precipitação acumulada e temperatura média.",
    )
    add_table(
        document,
        ["Cidade", "Dias", "Dias com chuva", "Taxa de chuva", "Precipitação total", "Temp. média"],
        city_summary_rows(city_summary),
        [1.1, 0.7, 1.1, 1.1, 1.3, 1.2],
    )
    for figure_path, caption in FIGURES:
        add_figure(document, figure_path, caption)

    add_heading(document, "4. Projeto e Modelagem do Artefato", 1)
    add_heading(document, "4.1 Requisitos do sistema", 2)
    add_bullet_list(
        document,
        [
            "Executar coleta parametrizada por .env.",
            "Validar dados antes de qualquer etapa analítica.",
            "Separar dados brutos, processados, modelos, relatórios e figuras.",
            "Registrar métricas e relatórios em arquivos auditáveis.",
            "Orquestrar atualização diária por Airflow.",
            "Manter interpretabilidade do modelo e rastreabilidade dos resultados.",
        ],
    )

    add_heading(document, "4.2 Arquitetura e tecnologias", 2)
    add_paragraph(
        document,
        "A arquitetura segue o fluxo: Open-Meteo API, camada raw, validação de "
        "contrato, camada processed, análise exploratória, treinamento, predições, "
        "relatórios, gráficos e orquestração diária por Airflow.",
    )
    add_bullet_list(
        document,
        [
            "Python para coleta, transformação, validação, EDA e modelagem.",
            "dotenv para carregamento de variáveis de ambiente.",
            ".venv para isolamento reprodutível do ambiente local.",
            "Airflow para agendamento diário às 05:00.",
            "CSV, JSON e PNG como artefatos auditáveis.",
        ],
    )

    add_heading(document, "4.3 Modelo lógico de dados", 2)
    add_table(
        document,
        ["Entidade", "Finalidade"],
        [
            ["weather_daily_observation", "Observação meteorológica diária por cidade."],
            ["rain_prediction_feature", "Linha supervisionada da cidade-alvo com variáveis e alvo do dia seguinte."],
            ["weather_city_comparison", "Resumo exploratório usado para comparação regional."],
            ["rain_model_report", "Métricas, matriz de confusão e interpretação do modelo."],
            ["airflow_dag_run", "Execução operacional diária do pipeline."],
        ],
        [2.2, 4.3],
    )

    add_heading(document, "5. Modelagem e Treinamento de Machine Learning", 1)
    add_heading(document, "5.1 Escolha e justificativa do algoritmo", 2)
    add_paragraph(
        document,
        "Foi escolhido um modelo probabilístico explicável baseado em sazonalidade "
        "mensal e persistência de chuva. Essa abordagem fortalece a defesa do "
        "projeto porque cada decisão pode ser rastreada até frequências históricas.",
    )

    add_heading(document, "5.2 Pipeline de treinamento", 2)
    add_numbered_list(
        document,
        [
            "Coletar dados atualizados das três cidades.",
            "Validar contrato dos dados brutos multi-cidade.",
            "Construir features da cidade-alvo Brasília.",
            "Validar contrato da base supervisionada.",
            "Treinar o modelo com histórico anterior ao corte temporal.",
            "Avaliar dados posteriores ao corte temporal.",
            "Publicar métricas, predições e gráficos.",
        ],
    )

    add_heading(document, "5.3 Métricas de avaliação", 2)
    add_paragraph(
        document,
        "Foram usadas acurácia, precisão, recall e F1-score. Como o objetivo "
        "operacional é alertar sobre chuva, recall e F1-score ajudam a avaliar "
        "o equilíbrio entre perda de dias chuvosos e alertas excedentes.",
    )

    add_heading(document, "6. Avaliação e Validação do Artefato", 1)
    add_heading(document, "6.1 Resultados obtidos", 2)
    test_metrics = backtest["test_metrics"]
    add_table(
        document,
        ["Métrica", "Valor no teste temporal"],
        [
            ["Acurácia", metric(test_metrics, "accuracy")],
            ["Precisão", metric(test_metrics, "precision")],
            ["Recall", metric(test_metrics, "recall")],
            ["F1-score", metric(test_metrics, "f1")],
        ],
        [2.4, 4.1],
    )
    matrix = test_metrics["confusion_matrix"]
    add_table(
        document,
        ["Classe", "Quantidade"],
        [
            ["Verdadeiro positivo", str(matrix["true_positive"])],
            ["Verdadeiro negativo", str(matrix["true_negative"])],
            ["Falso positivo", str(matrix["false_positive"])],
            ["Falso negativo", str(matrix["false_negative"])],
        ],
        [2.8, 3.7],
    )

    add_heading(document, "6.2 Análise crítica", 2)
    add_paragraph(
        document,
        f"O teste temporal apresentou F1-score de {metric(test_metrics, 'f1')}. "
        "O resultado indica que o mecanismo captura parte relevante dos dias "
        "chuvosos e mantém interpretação direta. A comparação com Goiânia e São "
        "Paulo fortalece o contexto, pois mostra que Brasília possui sazonalidade "
        "parecida com Goiânia e comportamento diferente de São Paulo em alguns meses.",
    )

    add_heading(document, "6.3 Recorte que precisa de validação", 2)
    add_paragraph(
        document,
        "O recorte que deve ser validado na apresentação é: Brasília como cidade-alvo "
        "da predição, Goiânia e São Paulo como cidades de comparação exploratória, "
        "chuva definida por precipitação diária >= 1 mm, atualização diária às 05:00 "
        "e corte temporal de avaliação em 2025-01-01.",
    )

    add_heading(document, "7. Conclusão", 1)
    add_heading(document, "7.1 Contribuições e impactos", 2)
    add_paragraph(
        document,
        "O projeto entrega uma solução robusta, reprodutível e auditável para "
        "previsão de chuva no dia seguinte em Brasília. A entrega cobre coleta "
        "automatizada, governança por contrato, comparação regional, gráficos de "
        "EDA, modelo explicável, Airflow e organização dos anexos técnicos.",
    )

    add_heading(document, "7.2 Trabalhos futuros", 2)
    add_numbered_list(
        document,
        [
            "Incluir baseline estatístico formal para comparação com o modelo atual.",
            "Testar outros algoritmos interpretáveis.",
            "Avaliar limiares alternativos de chuva.",
            "Adicionar painel operacional para consulta das predições.",
            "Ampliar a comparação para outras capitais do Centro-Oeste e Sudeste.",
        ],
    )

    add_heading(document, "8. Referências bibliográficas", 1)
    add_paragraph(
        document,
        "PEFFERS, Ken et al. A Design Science Research Methodology for Information "
        "Systems Research. Journal of Management Information Systems, v. 24, n. 3, "
        "p. 45-77, 2007.",
    )
    add_paragraph(
        document,
        "OPEN-METEO. Historical Weather API. Disponível em: "
        "https://open-meteo.com/en/docs/historical-weather-api.",
    )
    add_paragraph(
        document,
        "APACHE AIRFLOW. DAGs documentation. Disponível em: "
        "https://airflow.apache.org/docs/.",
    )

    add_heading(document, "9. Anexos", 1)
    add_heading(document, "ANEXO I - Termo de abertura de projeto", 2)
    add_table(
        document,
        ["Campo", "Valor"],
        [
            ["Projeto", PROJECT_NAME],
            ["Estudante", STUDENT_NAME],
            ["Matrícula", STUDENT_ID],
            ["Cidade-alvo", "Brasília"],
            ["Cidades de comparação", "Goiânia e São Paulo"],
            ["Artefato", "Pipeline robusto de dados, EDA, modelo e orquestração diária."],
        ],
        [1.8, 4.7],
    )

    add_heading(document, "ANEXO II - Arquivos técnicos para entrega", 2)
    add_table(
        document,
        ["Anexo", "Arquivo sugerido"],
        [
            ["Guia de execução", "docs/GUIA_EXECUCAO_AIRFLOW_E_TESTES.md"],
            ["Docker Compose Airflow", "docker-compose.airflow.yml"],
            ["Configuração de ambiente", ".env.example"],
            ["Configuração das cidades", "config/weather_locations.json"],
            ["Contratos de dados", "config/data_contract_weather_daily.json e config/data_contract_weather_features.json"],
            ["DAG Airflow", "airflow/dags/predicao_chuvas_brasilia_daily.py"],
            ["Execução diária", "outputs/reports/daily_pipeline_run_report.json"],
            ["Resumo comparativo", "outputs/reports/weather_city_comparison_summary.csv"],
            ["Métricas do modelo", "outputs/reports/rain_model_backtest.json"],
            ["Predições", "outputs/reports/rain_predictions.csv"],
            ["Gráficos EDA", "outputs/figures/*.png"],
            ["Código-fonte", "src/projeto_integrador/"],
        ],
        [2.0, 4.5],
    )


def read_city_summary(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8", newline="") as file:
        return list(csv.DictReader(file))


def city_summary_rows(summary: list[dict[str, str]]) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in summary:
        rows.append(
            [
                row["city_name"],
                row["day_count"],
                row["rainy_days"],
                percent(float(row["rain_rate"])),
                f"{float(row['total_precipitation_mm']):.1f} mm",
                f"{float(row['avg_temperature_c']):.1f} °C",
            ]
        )
    return rows


def summary_items() -> list[str]:
    return [
        "1. Introdução",
        "2. Metodologia Design Science Research (DSR)",
        "3. Engenharia e Análise de Dados",
        "4. Projeto e Modelagem do Artefato",
        "5. Modelagem e Treinamento de Machine Learning",
        "6. Avaliação e Validação do Artefato",
        "7. Conclusão",
        "8. Referências bibliográficas",
        "9. Anexos",
    ]


def add_centered_heading(document: Document, text: str, level: int) -> None:
    paragraph = add_heading(document, text, level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_heading(document: Document, text: str, level: int):
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_paragraph(document: Document, text: str):
    paragraph = document.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return paragraph


def add_bullet_list(document: Document, items: list[str]) -> None:
    numberings = get_numbering_ids(document)
    for item in items:
        paragraph = document.add_paragraph(item)
        apply_numbering(paragraph, numberings["bullet"])
        paragraph.paragraph_format.space_after = Pt(4)


def add_numbered_list(document: Document, items: list[str]) -> None:
    numberings = get_numbering_ids(document)
    for item in items:
        paragraph = document.add_paragraph(item)
        apply_numbering(paragraph, numberings["decimal"])
        paragraph.paragraph_format.space_after = Pt(4)


def add_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_inches: list[float],
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.autofit = False
    set_table_borders(table)
    set_table_width(table, widths_inches)

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        cell = header_cells[index]
        cell.text = header
        shade_cell(cell, "F2F4F7")
        set_cell_vertical_alignment(cell)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for row_values in rows:
        row_cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cell = row_cells[index]
            cell.text = value
            set_cell_vertical_alignment(cell)

    set_table_width(table, widths_inches)
    document.add_paragraph("")


def add_figure(document: Document, path: Path, caption: str) -> None:
    if not path.exists():
        add_paragraph(document, f"{caption} Arquivo não encontrado: {path.name}")
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.3))
    caption_paragraph = document.add_paragraph(caption)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.runs[0].italic = True


def get_numbering_ids(document: Document) -> dict[str, int]:
    cache_key = id(document)
    if cache_key not in NUMBERING_CACHE:
        NUMBERING_CACHE[cache_key] = create_numbering_definitions(document)
    return NUMBERING_CACHE[cache_key]


def create_numbering_definitions(document: Document) -> dict[str, int]:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
    ]
    next_abstract_id = max(abstract_ids, default=0) + 1
    next_num_id = max(num_ids, default=0) + 1

    bullet_abstract = add_abstract_numbering(
        numbering,
        abstract_id=next_abstract_id,
        num_format="bullet",
        level_text="\u2022",
    )
    bullet_num_id = add_numbering_instance(numbering, next_num_id, bullet_abstract)

    decimal_abstract = add_abstract_numbering(
        numbering,
        abstract_id=next_abstract_id + 1,
        num_format="decimal",
        level_text="%1.",
    )
    decimal_num_id = add_numbering_instance(numbering, next_num_id + 1, decimal_abstract)

    return {"bullet": bullet_num_id, "decimal": decimal_num_id}


def add_abstract_numbering(numbering, abstract_id: int, num_format: str, level_text: str) -> int:
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)

    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), num_format)
    level.append(fmt)

    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), level_text)
    level.append(text)

    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)

    ppr = OxmlElement("w:pPr")
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    ppr.append(indent)
    level.append(ppr)

    abstract.append(level)
    numbering.append(abstract)
    return abstract_id


def add_numbering_instance(numbering, num_id: int, abstract_id: int) -> int:
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))

    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)

    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")

    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_pr.append(ilvl)

    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.append(num_id_element)

    ppr.append(num_pr)


def set_table_width(table, widths_inches: list[float]) -> None:
    for row in table.rows:
        for index, width in enumerate(widths_inches):
            row.cells[index].width = Inches(width)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "BFBFBF")


def shade_cell(cell, fill: str) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell_properties.append(shading)


def set_cell_vertical_alignment(cell) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_text(paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def metric(metrics: dict[str, object], key: str) -> str:
    return f"{float(metrics[key]):.4f}".replace(".", ",")


def percent(value: float) -> str:
    return f"{value * 100:.1f}%".replace(".", ",")


if __name__ == "__main__":
    main()
